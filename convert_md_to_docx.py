# -*- coding: utf-8 -*-
"""
마크다운 문서를 DOCX(Word) 문서로 변환하는 스크립트입니다.
BeautifulSoup과 python-docx를 사용하여 문서 스타일을 맞춤으로 깔끔하게 설정합니다.
"""

import os
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- XML 조작 함수 (테이블 스타일링용) ---

def set_cell_background(cell, hex_color):
    """셀 배경색 설정"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """셀 안쪽 여백(패딩) 설정 (단위: dxa, 20 dxa = 1 pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, hex_color="2E75B6", size="36"):
    """셀의 왼쪽 테두리만 설정하고 나머지는 없앰 (인용구 박스용)"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    borders.append(left)
    
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        borders.append(b)
        
    tcPr.append(borders)

def set_table_borders(table, hex_color="CCCCCC"):
    """테이블의 상하 테두리 및 내부 수평 테두리만 설정 (클린 스타일)"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

# --- 텍스트 렌더링 헬퍼 ---

def apply_run_font(run, font_name="맑은 고딕", size_pt=10.5, color_rgb=None, bold=False, italic=False):
    """Run에 폰트 스타일을 적용"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    # 한글 폰트 적용을 위해 동아시아 폰트 명시적 지정
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def add_html_run(p, element, font_name="맑은 고딕", size_pt=10.5, default_color=None):
    """HTML 요소 내부의 스타일(strong, em 등)을 분석하여 문단 p에 Run으로 추가"""
    for child in element.children:
        if isinstance(child, str):
            if child.strip() == "" and len(list(element.children)) > 1:
                # 불필요한 공백 문자 방지
                continue
            run = p.add_run(child)
            apply_run_font(run, font_name, size_pt, color_rgb=default_color)
        elif child.name in ['strong', 'b']:
            text = child.get_text()
            run = p.add_run(text)
            apply_run_font(run, font_name, size_pt, color_rgb=default_color, bold=True)
        elif child.name in ['em', 'i']:
            text = child.get_text()
            run = p.add_run(text)
            apply_run_font(run, font_name, size_pt, color_rgb=default_color, italic=True)
        elif child.name == 'a':
            text = child.get_text()
            run = p.add_run(text)
            apply_run_font(run, font_name, size_pt, color_rgb=RGBColor(0, 51, 153), bold=True)
            run.underline = True
        elif child.name == 'code':
            text = child.get_text()
            run = p.add_run(text)
            apply_run_font(run, "Consolas", size_pt - 0.5, color_rgb=RGBColor(199, 37, 78))
        else:
            # 재귀 처리
            add_html_run(p, child, font_name, size_pt, default_color)

# --- 마크다운 요소별 핸들러 ---

def handle_blockquote(doc, bq_element):
    """인용 블록(Callout) 처리: 1행 1열 테이블로 옅은 파란색 배경과 왼쪽 굵은 테두리 적용"""
    text = bq_element.get_text().strip()
    
    # GitHub 스타일 경고 [!NOTE] 태그 제거 및 감지
    is_note = False
    if text.startswith('[!NOTE]'):
        is_note = True
        for p in bq_element.find_all('p'):
            p_text = p.get_text().strip()
            if p_text.startswith('[!NOTE]'):
                # 태그만 제거하거나 [!NOTE] 텍스트가 단독으로 있는 문단 삭제
                if p_text == '[!NOTE]':
                    p.decompose()
                else:
                    # [!NOTE] 부분을 텍스트에서 삭제
                    new_text = p_text.replace('[!NOTE]', '').strip()
                    p.string = new_text
                    
    # 테이블 생성 (1x1)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 가로 크기 지정 (약 6.2인치)
    table.columns[0].width = Inches(6.2)
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    
    # 옅은 배경색 및 안쪽 패딩 설정
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    # 왼쪽 파란색 강조선 테두리 설정
    set_cell_left_border(cell, hex_color="1F77B4", size="36")
    
    first_p = True
    for child in bq_element.children:
        if child.name in ['p', 'ul', 'ol']:
            if child.name == 'p':
                if first_p:
                    p = cell.paragraphs[0]
                    first_p = False
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                add_html_run(p, child, size_pt=10, default_color=RGBColor(51, 51, 51))
                
            elif child.name in ['ul', 'ol']:
                for li in child.find_all('li'):
                    if first_p:
                        p = cell.paragraphs[0]
                        first_p = False
                    else:
                        p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                    
                    # 불릿 포인트 수동 추가
                    bullet_run = p.add_run("• ")
                    apply_run_font(bullet_run, "맑은 고딕", 10, bold=True, color_rgb=RGBColor(31, 119, 180))
                    add_html_run(p, li, size_pt=10, default_color=RGBColor(51, 51, 51))

def handle_image(doc, img_element, md_dir):
    """이미지 삽입 및 캡션 추가"""
    src = img_element.get('src')
    alt = img_element.get('alt', '이미지')
    
    # 실행 경로 기준 이미지 절대 경로 구하기
    img_path = os.path.normpath(os.path.join(md_dir, src))
    
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        # 문서에 이미지 삽입 (가로 크기 5.5인치로 고정 조정)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        
        # 이미지 아래 캡션(설명) 삽입
        if alt:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(12)
            
            run_cap = p_cap.add_run(f"▲ {alt}")
            apply_run_font(run_cap, "맑은 고딕", 9, color_rgb=RGBColor(128, 128, 128), italic=True)
    else:
        # 파일이 없을 시 안내 문구 삽입
        p = doc.add_paragraph()
        run = p.add_run(f"[이미지를 찾을 수 없음: {src} ({alt})]")
        apply_run_font(run, "맑은 고딕", 10, color_rgb=RGBColor(204, 0, 0))

def handle_table(doc, table_element):
    """표 데이터 구성 및 스타일링"""
    rows = table_element.find_all('tr')
    if not rows:
        return
        
    cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 6.2인치 가로 공간을 균등 분할
    col_width = Inches(6.2 / cols)
    
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        row_cells = table.add_row().cells
        
        # 첫 번째 행은 헤더로 지정
        is_header = (row_idx == 0) or any(cell.name == 'th' for cell in cells)
        
        for col_idx, cell in enumerate(cells):
            if col_idx < len(row_cells):
                docx_cell = row_cells[col_idx]
                docx_cell.width = col_width
                
                # 셀 안쪽 여백 지정
                set_cell_margins(docx_cell, top=100, bottom=100, left=120, right=120)
                
                p = docx_cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.1
                
                cell_text = cell.get_text().strip()
                
                # 정렬 설정: 헤더는 가운데 정렬, 숫자는 오른쪽 정렬, 나머지는 왼쪽 정렬
                if is_header:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_background(docx_cell, "EAEFF5") # 파란빛 도는 밝은 회색
                    run = p.add_run(cell_text)
                    apply_run_font(run, "맑은 고딕", 10, color_rgb=RGBColor(46, 117, 182), bold=True)
                else:
                    # 숫자가 포함된 수치 데이터인지 판별 (예: 25,560, 8.4%, 2,701 등)
                    is_numeric = re.match(r'^[\d,.\-%]+(원|건|분|초|%|점)?$', cell_text) or cell_text.replace(',', '').replace('.', '').replace('-', '').isdigit()
                    if is_numeric:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    add_html_run(p, cell, size_pt=9.5)
                    
    set_table_borders(table, hex_color="CCCCCC")
    
    # 표 뒤에 약간의 여백 추가를 위한 빈 단락
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(8)

def handle_hr(doc):
    """구분선 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    # 단락 하단에 연한 회색 테두리 적용
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D0D0"/></w:pBdr>')
    pPr.append(pBdr)

# --- 메인 변환 로직 ---

def convert_md_to_docx(md_filepath, docx_filepath):
    """마크다운 파일을 읽고 파싱하여 깔끔하게 스타일링된 DOCX로 저장합니다."""
    print(f"읽는 중: {md_filepath}")
    with open(md_filepath, 'r', encoding='utf-8') as f:
        md_text = f.read()
        
    # 마크다운 -> HTML 변환 (테이블 및 fenced_code 지원)
    html_text = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html_text, 'html.parser')
    
    md_dir = os.path.dirname(md_filepath)
    
    # DOCX 문서 생성
    doc = Document()
    
    # 페이지 설정 (US Letter 크기, 1인치 여백)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # HTML 최상위 요소 순회하며 처리
    for element in soup.children:
        if element.name is None:
            continue
            
        if element.name == 'h1':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(element.get_text())
            # Heading 1: 크기 18pt, 볼드, 진한 파란색
            apply_run_font(run, "맑은 고딕", 18, color_rgb=RGBColor(31, 78, 121), bold=True)
            
            # 제목 밑에 구분선 추가
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1F4E79"/></w:pBdr>')
            pPr.append(pBdr)
            
        elif element.name == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(element.get_text())
            # Heading 2: 크기 13pt, 볼드, 남색
            apply_run_font(run, "맑은 고딕", 13, color_rgb=RGBColor(46, 117, 182), bold=True)
            
        elif element.name == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(element.get_text())
            # Heading 3: 크기 11.5pt, 볼드, 짙은 회색
            apply_run_font(run, "맑은 고딕", 11.5, color_rgb=RGBColor(89, 89, 89), bold=True)
            
        elif element.name == 'p':
            # 내부에 이미지가 단독으로 있는지 확인
            img_tag = element.find('img')
            if img_tag:
                handle_image(doc, img_tag, md_dir)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_html_run(p, element)
                
        elif element.name == 'blockquote':
            handle_blockquote(doc, element)
            
        elif element.name == 'table':
            handle_table(doc, element)
            
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                # 불릿 추가
                bullet_run = p.add_run("• ")
                apply_run_font(bullet_run, "맑은 고딕", 10.5, bold=True, color_rgb=RGBColor(46, 117, 182))
                add_html_run(p, li)
                
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li'), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                # 숫자 인덱스 추가
                num_run = p.add_run(f"{idx}. ")
                apply_run_font(num_run, "맑은 고딕", 10.5, bold=True, color_rgb=RGBColor(46, 117, 182))
                add_html_run(p, li)
                
        elif element.name == 'hr':
            handle_hr(doc)
            
        elif element.name in ['pre', 'code']:
            # 코드 블록 처리 (필요시)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(element.get_text())
            apply_run_font(run, "Consolas", 9.5, color_rgb=RGBColor(51, 51, 51))
            
    # 문서 저장
    print(f"저장 중: {docx_filepath}")
    doc.save(docx_filepath)
    print("변환 완료!")

if __name__ == "__main__":
    # 상대경로 지정
    src_md = os.path.join("yes24", "docs", "eda_report.md")
    dest_docx = os.path.join("yes24", "docs", "eda_report.docx")
    
    # 디렉토리 존재 확인 후 변환
    if os.path.exists(src_md):
        convert_md_to_docx(src_md, dest_docx)
    else:
        print(f"오류: 마크다운 파일이 존재하지 않습니다: {src_md}")
